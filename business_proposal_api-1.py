from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi.responses import FileResponse
from docx import Document
import os
import requests
from googletrans import Translator


app = FastAPI()

# Allow CORS for frontend communication
origins = [
    "http://localhost",
    "http://localhost:3000",  # Frontend development environment
    "*"  # Allow all origins for testing purposes; restrict in production
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define required topics
required_topics = [
    "Company Overview", "Mission and Vision Statement", "Executive Summary",
    "Owners and Partnerships", "Industry Overview and Trends", "Competition",
    "Problem Statement", "Marketing Plan", "Proposed Solution", "Market Analysis",
    "Sustainable Practices", "Implementation Timeline", "Staff Names",
    "Financial Objectives", "Exit Strategy", "Conclusion"
]

# Define proposal templates
templates = {
    "Template 1": [
        "Company Overview", "Mission and Vision Statement", "Executive Summary",
        "Problem Statement", "Proposed Solution", "Market Analysis",
        "Marketing Plan", "Implementation Timeline", "Conclusion"
    ],
    "Template 2": [
        "Executive Summary", "Problem Statement", "Market Analysis",
        "Proposed Solution", "Marketing Plan", "Implementation Timeline",
        "Conclusion", "Company Overview", "Mission and Vision Statement"
    ],
    "Template 3": [
        "Market Analysis", "Problem Statement", "Proposed Solution",
        "Marketing Plan", "Implementation Timeline", "Conclusion",
        "Mission and Vision Statement", "Company Overview", "Executive Summary"
    ]
}

# OCR.space API details
OCR_SPACE_API_URL = "https://api.ocr.space/parse/image"
OCR_SPACE_API_KEY = "K81609764688957"  # Replace with your OCR.space API key

# Groq API details
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = "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"  # Replace with your Groq API key

def extract_text_from_pdf(file: UploadFile):
    response = requests.post(
        OCR_SPACE_API_URL,
        files={'file': (file.filename, file.file)},
        data={'apikey': OCR_SPACE_API_KEY, 'language': 'eng'}
    )
    result = response.json()
    return result.get('ParsedResults', [{}])[0].get('ParsedText', '')

def identify_missing_topics(extracted_text):
    missing = []
    for topic in required_topics:
        if topic.lower() not in extracted_text.lower():
            missing.append(topic)
    return missing

def sinhala_to_english_translation(input_text):
    translator = Translator()
    translation = translator.translate(input_text, src='si', dest='en')
    return translation.text


@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a PDF file, extract its content, and identify missing topics.
    Dynamically forward to endpoints for missing topics and redirect to proposal generation.
    """
    extracted_text = extract_text_from_pdf(file)
    missing_topics = identify_missing_topics(extracted_text)
    
    # Collect inputs for missing topics
    collected_inputs = {}
    for topic in missing_topics:
        endpoint_url = f"http://localhost:8000/missing/{topic.replace(' ', '-').lower()}"
        # Simulate getting inputs for missing topics. In production, use a frontend or user interface to get details.
        sample_input = {"details": f"Sample details for {topic}"}
        response = requests.post(endpoint_url, json=sample_input)
        if response.status_code == 200:
            collected_inputs[topic] = response.json().get("details")

    # Prepare metadata and inputs for proposal generation
    metadata = {
        "name": "Sample Business Name",
        "domain": "Sample Domain",
        "is_existing": "Yes",
        "user_instructions": "Sample user instructions"
    }

    # Redirect to proposal generation with all data
    redirect_response = requests.post(
        "http://localhost:8000/generate-proposal/",
        data={
            "business_name": metadata["name"],
            "business_domain": metadata["domain"],
            "is_existing": metadata["is_existing"],
            "user_instructions": metadata["user_instructions"],
            "selected_template": "Template 1"  # Or dynamically selected
        },
        files={"file": file.file}
    )
    return {"status": "Proposal generated", "proposal_url": redirect_response.url}


# Models for user input
class TopicInput(BaseModel):
    details: str
    


# Example separate endpoints for each topic
@app.post("/missing/{topic}")
async def collect_missing_topic(topic: str, input: TopicInput):
    """
    Dynamic endpoint for providing information for any missing topic.
    The topic name is passed dynamically via the URL path.
    """
    topic_mapping = {
        "company-overview": "Company Overview",
        "mission-vision": "Mission and Vision Statement",
        "executive-summary": "Executive Summary",
        "owners-and-partnerships": "Owners and Partnerships",
        "industry-overview-and-trends": "Industry Overview and Trends",
        "competition": "Competition",
        "problem-statement": "Problem Statement",
        "marketing-plan": "Marketing Plan",
        "proposed-solution": "Proposed Solution",
        "market-analysis": "Market Analysis",
        "sustainable-practices": "Sustainable Practices",
        "implementation-timeline": "Implementation Timeline",
        "staff-names": "Staff Names",
        "financial-objectives": "Financial Objectives",
        "exit-strategy": "Exit Strategy",
        "conclusion": "Conclusion"
    }

    # Resolve the topic name from the mapping
    resolved_topic = topic_mapping.get(topic.lower(), topic.replace("-", " ").title())

    return {"topic": resolved_topic, "details": input.details}



# Proposal Generation Endpoint
def generate_proposal_with_groq(metadata, user_inputs, topic_order):
    content = {}
    for topic in topic_order:
        details = user_inputs.get(topic, "")
        prompt = (
            f"Create detailed content for the topic '{topic}'.\n"
            f"Business Name: {metadata['name']}\n"
            f"Domain: {metadata['domain']}\n"
            f"Existing Business: {metadata['is_existing']}\n"
            f"Instructions: {metadata['user_instructions']}\n"
            f"Details: {details}\n"
        )

        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={'Authorization': f'Bearer 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'},
            json={
                "messages": [{"role": "user", "content": prompt}],
                "model": "llama-3.1-8b-instant"
            }
        )
        result = response.json()
        topic_content = result['choices'][0]['message']['content']
        content[topic] = topic_content

    return content

def save_proposal_to_word(content, output_path):
    doc = Document()
    doc.add_heading("Business Proposal", level=0)
    for topic, details in content.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(details)
        doc.add_page_break()
    doc.save(output_path)

# Add global variable to store user instructions temporarily
user_instructions_storage = {"instructions": ""}

@app.post("/submit-instructions/")
async def submit_instructions(
    instructions: str = Form(...),
    language: str = Form("en")
):
    """
    Endpoint to accept user instructions in English or Sinhala.
    Automatically detects and translates instructions if necessary.
    """
    global user_instructions_storage  # Use the global storage
    translator = Translator()

    # Detect language if not provided
    detected_lang = translator.detect(instructions).lang if language == "auto" else language

    if detected_lang == 'si':
        print("Detected Sinhala input. Translating to English...")
        instructions = sinhala_to_english_translation(instructions)

    # Store the processed instructions
    user_instructions_storage["instructions"] = instructions

    return {
        "original_language": detected_lang,
        "instructions": instructions,
        "message": "User instructions processed and stored successfully."
    }


@app.post("/submit-voice-instructions/")
async def submit_voice_instructions(
    voice_instructions: str = Form(...),
    language: str = Form("en")
):
    """
    Endpoint to accept user instructions converted from voice to text in English or Sinhala.
    Automatically detects and translates instructions if necessary.
    """
    global user_instructions_storage  # Use the global storage
    translator = Translator()

    # Detect language if not provided
    detected_lang = translator.detect(voice_instructions).lang if language == "auto" else language

    if detected_lang == 'si':
        print("Detected Sinhala input in voice instructions. Translating to English...")
        voice_instructions = sinhala_to_english_translation(voice_instructions)

    # Store the processed voice instructions
    user_instructions_storage["instructions"] = voice_instructions

    return {
        "original_language": detected_lang,
        "voice_instructions": voice_instructions,
        "message": "Voice instructions processed and stored successfully."
    }


@app.post("/generate-proposal/")
async def generate_proposal(
    file: UploadFile = File(...),
    business_name: str = Form(...),
    business_domain: str = Form(...),
    is_existing: str = Form(...),
    user_instructions: str = Form(None),  # Make this optional
    selected_template: str = Form(...),
):
    """
    Endpoint to generate the business proposal. Pulls user instructions from 
    global storage if not provided explicitly.
    """
    global user_instructions_storage  # Use the global storage

    # Use instructions from request if provided, else fallback to stored instructions
    if not user_instructions:
        user_instructions = user_instructions_storage["instructions"]

    if not user_instructions:
        return {"error": "User instructions are missing. Please provide them or submit via the instructions endpoint."}

    print(f"Business Name: {business_name}")
    print(f"User Instructions: {user_instructions}")

    # Translate Sinhala input if detected
    translator = Translator()
    detected_lang = translator.detect(user_instructions).lang
    if detected_lang == 'si':
        print("Detected Sinhala input for user instructions. Translating to English...")
        user_instructions = sinhala_to_english_translation(user_instructions)

    # Step 1: Extract text from PDF
    extracted_text = extract_text_from_pdf(file)

    # Step 2: Identify missing topics
    missing_topics = identify_missing_topics(extracted_text)

    # Step 3: Collect dummy user inputs for missing topics (can be replaced with real inputs)
    user_inputs = {topic: "Details not provided" for topic in missing_topics}

    # Step 4: Collect business metadata
    metadata = {
        "name": business_name,
        "domain": business_domain,
        "is_existing": is_existing,
        "user_instructions": user_instructions,
    }

    # Step 5: Determine topic order based on selected template
    topic_order = templates[selected_template]

    # Step 6: Generate business proposal using Groq API
    proposal_content = generate_proposal_with_groq(metadata, user_inputs, topic_order)

    # Step 7: Save proposal to Word document
    output_path = os.path.join(os.getcwd(), "Business_Proposal.docx")
    save_proposal_to_word(proposal_content, output_path)

    # Return the generated document as a downloadable file
    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="Business_Proposal.docx")
